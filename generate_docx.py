#!/usr/bin/env python3
"""
Generate Finistral AI research paper as .docx following GR Journals Submission Template.

Formatting specs (from template):
  - Title: Calibri 20pt, Bold, Title Case
  - Author: TNR 10pt
  - Affiliation: TNR 9pt, Italic
  - Email: TNR 9pt
  - Abstract heading: Calibri 12pt, Bold
  - Abstract body: Calibri 10.5pt, single spacing
  - Keywords: TNR 9pt, Italic
  - Section headings: TNR 10.5pt, Bold, numbered
  - Body text: TNR 10.5pt, single spacing, justified
  - CRediT: TNR Bold, 10.5pt, paragraph spacing before 12pt after 0pt
  - References: TNR 10.5pt, single, justified
  - Figure captions: "Fig." in bold
  - Page: A4, 1-inch margins
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

FIGURES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "figures")
OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "Finistral_AI_GRJournals.docx")

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_run(paragraph, text, font_name="Times New Roman", font_size=10.5,
            bold=False, italic=False, superscript=False, color=None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if superscript:
        run.font.superscript = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_hyperlink(paragraph, text, url, font_name="Times New Roman", font_size=10.5):
    """Add a blue hyperlink to a paragraph."""
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.underline = True
    return run

def body_paragraph(doc, text_parts, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   space_after=Pt(0), space_before=Pt(0)):
    """
    Create a body-text paragraph.
    text_parts is a list of dicts: {text, bold, italic, superscript, font_name, font_size, color}
    """
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    p.paragraph_format.line_spacing = 1.0
    for tp in text_parts:
        add_run(p,
                tp.get("text", ""),
                font_name=tp.get("font_name", "Times New Roman"),
                font_size=tp.get("font_size", 10.5),
                bold=tp.get("bold", False),
                italic=tp.get("italic", False),
                superscript=tp.get("superscript", False),
                color=tp.get("color", None))
    return p

def simple_body(doc, text, space_after=Pt(6)):
    return body_paragraph(doc, [{"text": text}], space_after=space_after)

def bold_body(doc, text, space_after=Pt(6)):
    return body_paragraph(doc, [{"text": text, "bold": True}], space_after=space_after)

def section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, f"{number}. {title}", bold=True, font_size=10.5)
    return p

def subsection_heading(doc, number, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, f"{number} {title}", bold=True, font_size=10.5)
    return p

def unnumbered_heading(doc, title, font_name="Times New Roman", font_size=10.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, title, bold=True, font_size=font_size, font_name=font_name)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, h, bold=True, font_size=10.5)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bold = False
            if isinstance(val, tuple):
                val, bold = val
            add_run(p, str(val), bold=bold, font_size=10.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table

def add_figure(doc, image_path, width_inches, caption_text, fig_num):
    """Add a figure with caption in 'Fig. X' format (bold label)."""
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Inches(width_inches))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    cap.paragraph_format.line_spacing = 1.0
    add_run(cap, f"Fig. {fig_num}. ", bold=True, font_size=9)
    add_run(cap, caption_text, font_size=9)

def add_reference(doc, num, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, f"[{num}] ", font_size=10.5)
    add_run(p, text, font_size=10.5)


# ── MAIN ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # ── Page Setup (A4, 1-inch margins) ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(0)

    # ══════════════════════════════════════════════════════════════════════════
    # TITLE BLOCK
    # ══════════════════════════════════════════════════════════════════════════

    # Title: Calibri 20pt, Bold
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "Finistral AI: Financial Sentiment Analyst",
            font_name="Calibri", font_size=20, bold=True)

    # Authors
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, "Zainab Mirza", font_size=10)
    add_run(p, "1", font_size=10, superscript=True)
    add_run(p, ", Ayan Javeed Shaikh", font_size=10)
    add_run(p, "2,*", font_size=10, superscript=True)
    add_run(p, ", Fazal Jalil Parkar", font_size=10)
    add_run(p, "3", font_size=10, superscript=True)

    # Affiliations
    affils = [
        ("1", "University of Mumbai, Mumbai, Maharashtra, India"),
        ("2", "Indiana University Bloomington, Bloomington, IN 47405, USA"),
        ("3", "Tata Consultancy Services (TCS), India"),
    ]
    for num, text in affils:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, num, font_size=9, italic=True, superscript=True)
        add_run(p, text, font_size=9, italic=True)

    # Email
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, "*", font_size=9, superscript=True)
    add_run(p, "Email: ayshaikh@iu.edu", font_size=9)

    # ══════════════════════════════════════════════════════════════════════════
    # ABSTRACT
    # ══════════════════════════════════════════════════════════════════════════

    # Abstract heading: Calibri 12pt, Bold
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Abstract", font_name="Calibri", font_size=12, bold=True)

    # Abstract text: Calibri 10.5pt
    abstract_text = (
        "We present Finistral-7B-LoRA, a parameter-efficient financial-sentiment classifier "
        "produced by Low-Rank Adapting the 7-billion-parameter Mistral-7B-v0.1 language model. "
        "The adapter is trained on the open-source FinGPT/fingpt-sentiment-train corpus comprising "
        "approximately 77,000 labelled sentences drawn from press releases, wire headlines, analyst "
        "commentary, and Reddit finance threads. Training is conducted on Indiana University\u2019s "
        "Big Red 200 supercomputer using two NVIDIA A100 (40 GB) GPUs, completing in under 120 "
        "minutes while updating only \u2248 0.2% of the model parameters. On the Financial PhraseBank "
        "benchmark (sentences_allagree split), Finistral-7B-LoRA achieves 99.56% accuracy and a "
        "weighted F1 score of 0.996, reducing test errors to just ten instances and outperforming "
        "five strong FinGPT baselines as well as FinBERT by substantial margins. We further "
        "demonstrate that the full-precision adapter occupies only 83.9 MB and that inference latency "
        "remains below 100 ms on both GPU and CPU, making the model practical for real-time financial "
        "news monitoring. The released repository Ayansk11/Finistral-7B_lora on Hugging Face contains "
        "the full-precision adapter, a 4-bit GGML variant, tokenizer files, and a Gradio-based "
        "inference interface for reproducible research."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, abstract_text, font_name="Calibri", font_size=10.5)

    # Keywords: TNR 9pt, Italic
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, "Keywords: ", font_size=9, italic=True)
    add_run(p, "financial sentiment analysis, Low-Rank Adaptation (LoRA), parameter-efficient "
            "fine-tuning (PEFT), Mistral-7B, large language models, FinGPT, Financial PhraseBank",
            font_size=9, italic=True)

    # ══════════════════════════════════════════════════════════════════════════
    # GRAPHICAL ABSTRACT
    # ══════════════════════════════════════════════════════════════════════════

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "Graphical Abstract", font_name="Calibri", font_size=12, bold=True)

    # Side by side images (Word can't do true side-by-side easily, so use a 2-col table)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in tbl.rows[0].cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    img3 = os.path.join(FIGURES_DIR, "image3.png")
    img5 = os.path.join(FIGURES_DIR, "image5.png")
    if os.path.exists(img3):
        tbl.rows[0].cells[0].paragraphs[0].add_run().add_picture(img3, width=Inches(2.8))
    if os.path.exists(img5):
        tbl.rows[0].cells[1].paragraphs[0].add_run().add_picture(img5, width=Inches(2.8))
    # Remove table borders
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                '</w:tcBorders>')
            tcPr.append(tcBorders)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "Left: Finistral-7B-LoRA vs Mistral-7B-Base performance comparison on Financial "
            "PhraseBank. Right: Confusion matrix showing near-perfect classification across all "
            "three sentiment classes.", font_size=9)

    # ══════════════════════════════════════════════════════════════════════════
    # NOVELTY STATEMENT
    # ══════════════════════════════════════════════════════════════════════════

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, "Novelty Statement: ", font_size=9, italic=True, bold=False)
    add_run(p, "This work is the first to apply Low-Rank Adaptation to the compute-efficient "
            "Mistral-7B backbone for financial sentiment classification, achieving state-of-the-art "
            "accuracy (99.56%) on Financial PhraseBank while updating only 0.2% of the model "
            "parameters. The resulting adapter is publicly released to enable low-cost, reproducible "
            "FinNLP research.", font_size=9, italic=True)

    # ══════════════════════════════════════════════════════════════════════════
    # 1. INTRODUCTION
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, 1, "Introduction")

    simple_body(doc, "Financial markets digest a relentless stream of earnings releases, "
        "macro-economic announcements, social-media chatter, and regulatory filings. Traders and "
        "risk managers depend on sentiment analysis systems that can transform this unstructured "
        "text into structured signals\u2014typically the triplet positive, neutral, or "
        "negative\u2014within seconds. Conventional rule-based lexicons [1] and classical "
        "machine-learning pipelines capture obvious polarity cues, yet they struggle with the "
        "domain\u2019s specialised jargon (\u201cEPS miss\u201d), pragmatic hedges "
        "(\u201cguidance was raised but remains below expectations\u201d), and rapidly evolving "
        "slang from forums such as WallStreetBets.")

    simple_body(doc, "Large language models (LLMs) have recently set new performance bars on "
        "general-domain sentiment tasks, but two obstacles limit their usefulness in finance:")

    p = body_paragraph(doc, [
        {"text": "Domain mismatch: ", "bold": True},
        {"text": "Out-of-the-box LLMs misinterpret sector-specific expressions and implicit market context."}
    ])

    p = body_paragraph(doc, [
        {"text": "Compute cost: ", "bold": True},
        {"text": "Full fine-tuning of multi-billion-parameter models demands prohibitive GPU hours and "
         "storage, placing them out of reach for most academic or boutique-fund settings."}
    ])

    simple_body(doc, "Parameter-Efficient Fine-Tuning (PEFT) techniques such as Low-Rank Adaptation "
        "(LoRA) [5] address both issues by inserting small trainable matrices into each transformer "
        "layer while freezing the backbone. This yields two attractive properties: (i) the modified "
        "model can be trained on a single commodity GPU, and (ii) the resulting adapter file is "
        "lightweight, making deployment trivial.")

    # 1.1 Research Questions
    subsection_heading(doc, "1.1", "Research Questions")

    simple_body(doc, "I investigate whether a LoRA-based adaptation of the compute-efficient "
        "Mistral-7B decoder can provide state-of-the-art sentiment classification in finance:")

    p = body_paragraph(doc, [
        {"text": "Research Question 1: ", "bold": True},
        {"text": "Can a LoRA-tuned Mistral-7B surpass specialised encoders such as FinBERT and the "
         "latest FinGPT adapters on the Financial PhraseBank benchmark?"}
    ])

    p = body_paragraph(doc, [
        {"text": "Research Question 2: ", "bold": True},
        {"text": "Does the PEFT recipe retain training- and inference-time efficiency suitable for "
         "realistic academic or start-up budgets?"}
    ])

    # 1.2 Our Approach
    subsection_heading(doc, "1.2", "Our Approach")

    simple_body(doc, "To answer these questions, I built Finistral-7B-LoRA. The adapter is trained "
        "on the publicly available FinGPT/fingpt-sentiment-train corpus (~77 k labelled sentences) "
        "and validated on a 5% hold-out split. Training runs for four epochs on two A100 (40 GB) "
        "GPUs of Indiana University\u2019s Big Red 200 supercomputer, completing in under 120 "
        "minutes while updating only \u2248 0.2% of the model parameters. Validation loss bottomed "
        "out at 0.1008 after epoch 2 and began to rise thereafter, so that checkpoint is used for "
        "all downstream evaluations.")

    # 1.3 Contributions
    subsection_heading(doc, "1.3", "Contributions")

    simple_body(doc, "My work makes three key contributions:")

    p = body_paragraph(doc, [
        {"text": "Model release: ", "bold": True},
        {"text": "I publish Finistral-7B-LoRA along with full-precision (83.9 MB) and 4-bit GGML "
         "(168 MB) weights, tokeniser files, and inference scripts."}
    ])

    p = body_paragraph(doc, [
        {"text": "State-of-the-art performance: ", "bold": True},
        {"text": "The adapter achieves 99.56% accuracy and 0.996 weighted F1 on Financial PhraseBank "
         "(sentences_allagree), reducing test errors to just ten instances and outperforming five "
         "strong FinGPT baselines by large margins."}
    ])

    p = body_paragraph(doc, [
        {"text": "Efficiency analysis: ", "bold": True},
        {"text": "We demonstrate that high-quality financial sentiment models can be trained in < 1 "
         "GPU-hour per epoch and deployed with minimal memory overhead, opening the door to low-cost, "
         "reproducible FinNLP research."}
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # 2. RELATED WORK
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, 2, "Related Work")

    simple_body(doc, "Early attempts at financial sentiment analysis relied on domain-specific "
        "lexicons such as the Loughran\u2013McDonald word lists [1], which map accounting terms to "
        "positive or negative polarity. Although lexicon rules are transparent, their coverage is "
        "limited and they struggle with figurative or contextual language. Subsequent classical "
        "machine-learning pipelines trained support-vector or recurrent networks on Financial "
        "PhraseBank [2]\u2014a 4.8 k-sentence news corpus with crowd-sourced sentiment "
        "labels\u2014achieving modest gains but still underperforming human readers.")

    simple_body(doc, "The advent of pre-trained language models (PLMs) led to a sharp performance "
        "jump. BERT [3] introduced deep bidirectional pre-training for language understanding, and "
        "FinBERT [4] fine-tunes BERT on a large corpus of SEC filings, significantly outperforming "
        "earlier feature-engineering baselines on both opinion mining and aspect sentiment tasks. "
        "FinBERT, however, requires updating the full 110 M parameters and thus incurs substantial "
        "GPU cost for domain shifts.")

    simple_body(doc, "To reduce compute requirements, the community has embraced Parameter-Efficient "
        "Fine-Tuning (PEFT). LoRA [5] freezes backbone weights and injects small rank-decomposition "
        "matrices, cutting trainable parameters by orders of magnitude while matching full "
        "fine-tuning quality. QLoRA [6] extends this with 4-bit quantisation, enabling fine-tuning "
        "of 65B-parameter models on a single 48 GB GPU. Surveys [7] catalogue dozens of PEFT "
        "variants\u2014including adapters, prefix-tuning\u2014and highlight their importance for "
        "resource-constrained practitioners. Zhang et al. [9] demonstrate that instruction tuning "
        "general-purpose LLMs on financial sentiment data yields strong performance, particularly "
        "where numerical understanding and contextual comprehension are vital.")

    simple_body(doc, "Within finance, the open-source FinGPT project [8] provides data-centric "
        "pipelines and a zoo of LoRA adapters for Llama-, Falcon- and Bloom-based models. These "
        "checkpoints form strong baselines in the annual FinNLP shared tasks [10], which benchmark "
        "sentiment, headline classification and market-impact prediction. Yet, published FinGPT "
        "adapters still lag behind specialised encoders like FinBERT on the Financial PhraseBank "
        "test split.")

    simple_body(doc, "Parallel to algorithmic advances, model architectures have evolved toward "
        "compute efficiency. The recently released Mistral-7B [11] employs sliding-window attention "
        "and grouped query projection to halve inference cost while retaining GPT-3-level accuracy. "
        "Combining Mistral\u2019s lean decoder with LoRA thus promises a sweet-spot of speed and "
        "accuracy that prior financial studies have not fully explored.")

    simple_body(doc, "Our work situates itself at this intersection: we leverage the Mistral-7B "
        "backbone and the FinGPT sentiment corpus, apply LoRA for parameter-efficient adaptation, "
        "and report state-of-the-art results on Financial PhraseBank, surpassing both FinBERT and "
        "existing FinGPT adapters.")

    # ══════════════════════════════════════════════════════════════════════════
    # 3. PROBLEM DEFINITION AND DATASET
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, 3, "Problem Definition and Dataset")

    subsection_heading(doc, "3.1", "Task Formulation")

    simple_body(doc, "Given a single English news headline or short text snippet x, predict a "
        "sentiment label y \u2208 {negative, neutral, positive}. Performance is evaluated with "
        "Accuracy and Weighted F1\u2014the same metrics adopted by FinNLP shared-task benchmarks.")

    subsection_heading(doc, "3.2", "Training Corpus")

    simple_body(doc, "I fine-tuned on the open-source FinGPT/fingpt-sentiment-train dataset [8]. "
        "After deduplication the corpus contains 76,772 sentences gathered from press releases, "
        "wire headlines, analyst commentary, and Reddit finance threads.")

    simple_body(doc, "Using train_test_split(test_size=0.05, seed=42) we obtain:")

    # Table 1: Train-Test Split
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    add_run(p, "Table 1: ", bold=True, font_size=9)
    add_run(p, "Train\u2013Test Split", font_size=9)

    add_table(doc,
              ["Split", "Train", "Val"],
              [["# Sentences", "72,933", "3,839"],
               ["Positives (%)", "24.7", "24.9"],
               ["Neutral (%)", "50.8", "50.2"],
               ["Negative (%)", "24.5", "24.9"]])

    doc.add_paragraph()  # spacer

    subsection_heading(doc, "3.3", "Preprocessing")

    p = body_paragraph(doc, [
        {"text": "Pipeline Cleaning ", "bold": True},
        {"text": "\u2013 Unicode-normalise, remove HTML artefacts, collapse whitespace."}
    ])

    p = body_paragraph(doc, [
        {"text": "Tokenisation: ", "bold": True},
        {"text": "Use the native "},
        {"text": "Mistral-7B-v0.1", "italic": True},
        {"text": " 32 k BPE tokeniser."}
    ])

    p = body_paragraph(doc, [
        {"text": "Prompt templating: ", "bold": True},
        {"text": "Each sentence is wrapped exactly as in the training script:"}
    ])

    # Prompt template (indented italic)
    for line in ["Instruction: What is the sentiment of this news? Please choose an answer from {negative/neutral/positive}",
                 "Input: {text}",
                 "Answer:"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, line, italic=True, font_size=10.5)

    simple_body(doc, "This \u201cinstruction-style\u201d format consistently triggers single-token "
        "answers that simplify label mapping.")

    p = body_paragraph(doc, [
        {"text": "Label Mapping: ", "bold": True},
        {"text": "Generation is truncated at the first decoded token; a regex maps negative \u2192 0, "
         "neutral \u2192 1, positive \u2192 2."}
    ])

    # Table 2: Label Distribution
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    add_run(p, "Table 2: ", bold=True, font_size=9)
    add_run(p, "Label Distribution", font_size=9)

    add_table(doc,
              ["Class", "# Sentences", "Share (%)"],
              [["Neutral", "1391", "61.6"],
               ["Positives", "570", "25.2"],
               ["Negative", "303", "13.4"]])

    doc.add_paragraph()

    simple_body(doc, "FinGPT captures contemporary finance slang (e.g., \u201cYOLO\u201d, "
        "\u201ctendies\u201d) that legacy corpora miss.")

    p = body_paragraph(doc, [
        {"text": "Financial PhraseBank [2] remains the community\u2019s reference dataset, enabling "
         "direct comparison with FinBERT and all "},
        {"text": "FinGPT LoRA baselines", "bold": True},
        {"text": ": "},
        {"text": "FinGPT-mt-Llama2-7B-LoRA", "italic": True},
        {"text": ", "},
        {"text": "FinGPT-Llama-3-8B-LoRA", "italic": True},
        {"text": ", "},
        {"text": "FinGPT-Falcon-7B-LoRA", "italic": True},
        {"text": ", and "},
        {"text": "FinGPT-Bloom-7B1-LoRA", "italic": True},
        {"text": "."}
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # 4. METHODOLOGY
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, 4, "Methodology")

    simple_body(doc, "This section explains how I adapt the 7-billion-parameter Mistral-7B-v0.1 "
        "language model to three-way financial sentiment classification with minimal compute and "
        "memory overhead.")

    subsection_heading(doc, "4.1", "Base Model")

    simple_body(doc, "Mistral-7B-v0.1 is a decoder-only Transformer that combines grouped-query "
        "attention and sliding-window positional encodings to halve inference cost compared with "
        "comparably sized GPT-style models [11]. We utilise Unsloth [13], an open-source framework "
        "that applies custom Triton kernels to accelerate LoRA fine-tuning by up to 2\u00d7 while "
        "reducing VRAM consumption by approximately 70%. We load the backbone in 8-bit NF4 format "
        "(load_in_8bit=True in Transformers) so that a single 40 GB A100 can accommodate both "
        "weights and optimiser states.")

    subsection_heading(doc, "4.2", "LoRA Adaptation")

    simple_body(doc, "I employed Low-Rank Adaptation (LoRA) [5] to inject trainable "
        "rank-decomposition matrices into every attention-projection and feed-forward linear layer "
        "while leaving the original weights frozen.")

    # Table 3: LoRA Hyperparameters
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    add_run(p, "Table 3: ", bold=True, font_size=9)
    add_run(p, "LoRA Hyperparameters", font_size=9)

    add_table(doc,
              ["Hyperparameter", "Value"],
              [["LoRA rank r", "16"],
               ["Scaling \u03b1", "32"],
               ["LoRA dropout", "0.10"],
               ["Trainable parameters", "\u2248 9 M (0.2% of 7B)"],
               ["Adapter size", "83.9 MB"]])

    doc.add_paragraph()

    simple_body(doc, "A second adapter file (ggml-adapter-model.bin, 168 MB) is produced via 4-bit "
        "weight quantisation for CPU/edge deployments.")

    subsection_heading(doc, "4.3", "Training Procedure")

    simple_body(doc, "I used the Hugging Face Trainer API with the following settings.")

    # Table 4: Training Arguments
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    add_run(p, "Table 4: ", bold=True, font_size=9)
    add_run(p, "Training Arguments", font_size=9)

    add_table(doc,
              ["Setting", "Value"],
              [["Training data", "FinGPT/fingpt-sentiment-train (95% split)"],
               ["Validation data", "5% held-out split, seed 42"],
               ["Optimiser", "AdamW, \u03b2=(0.9, 0.999), \u03b5=1e-8"],
               ["Learning rate", "1 \u00d7 10\u207b\u2074"],
               ["Scheduler", "Cosine decay, 10 warm-up steps"],
               ["Epochs", "4"],
               ["Per-device batch", "32 (train & eval)"],
               ["Total batch (2\u00d7A100)", "64"],
               ["Precision", "bf16"],
               ["Gradient accumulation", "1"],
               ["Weight decay", "0.0"],
               ["Logging / eval", "every 100 steps"]])

    doc.add_paragraph()

    p = body_paragraph(doc, [
        {"text": "Training on two "},
        {"text": "A100 40 GB GPUs", "bold": True},
        {"text": " of Indiana University\u2019s "},
        {"text": "Big Red 200", "bold": True},
        {"text": " completes in \u2248 "},
        {"text": "120 minutes", "bold": True},
        {"text": " (4560 optimiser steps). Validation loss bottoms out at "},
        {"text": "0.1008 after epoch 2", "bold": True},
        {"text": "; that checkpoint is consequently chosen for all downstream experiments."}
    ])

    # Table 5: Loss Trajectory
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    add_run(p, "Table 5: ", bold=True, font_size=9)
    add_run(p, "Training Result \u2013 Validation-loss Trajectory (best at epoch 2)", font_size=9)

    add_table(doc,
              ["Epoch", "Step", "Training Loss", "Validation Loss"],
              [["1", "1140", "0.0680", "0.1121"],
               ["2", "2280", "0.1337", ("0.1009", True)],
               ["3", "3420", "0.0499", "0.11479"],
               ["4", "4560", "0.0014", "0.1599"]])

    doc.add_paragraph()

    subsection_heading(doc, "4.4", "Inference & Deployment")

    simple_body(doc, "At inference time we load the frozen backbone plus the 83.9 MB adapter in "
        "half-precision (torch.bfloat16) on GPU or in 4-bit GGML on CPU. A Gradio GUI wraps the "
        "pipeline, supporting sentence queries.")

    simple_body(doc, "The complete training recipe, adapters, and inference code with Gradio GUI are "
        "publicly available at Ayansk11/Finistral-7B_lora on Hugging Face for reproducible research.")

    # GUI Figures
    add_figure(doc, os.path.join(FIGURES_DIR, "image1.jpeg"), 4.5,
               "Finistral-7B-LoRA GUI", 1)
    add_figure(doc, os.path.join(FIGURES_DIR, "image2.jpeg"), 4.5,
               "Finistral-7B-LoRA GUI", 2)

    # ══════════════════════════════════════════════════════════════════════════
    # 5. EXPERIMENT
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, 5, "Experiment")

    simple_body(doc, "This section describes the evaluation protocol, baselines, metrics, and "
        "implementation details we use to assess Finistral-7B-LoRA on sentence-level financial "
        "sentiment classification.")

    subsection_heading(doc, "5.1", "Baselines")

    simple_body(doc, "All baselines are evaluated with the same instruction prompt, identical "
        "truncation rules, and greedy decoding.")

    # Table 6: Baselines
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    add_run(p, "Table 6: ", bold=True, font_size=9)
    add_run(p, "Baselines", font_size=9)

    add_table(doc,
              ["ID", "Model", "Parameters", "Fine-tuned?"],
              [["B\u2081", "Mistral-7B-Base", "7 B", "\u00d7"],
               ["B\u2082", "FinGPT-mt-Llama2-7B-LoRA", "7 B + 9 M", "\u2713"],
               ["B\u2083", "FinGPT-Llama-3-8B-LoRA", "8 B + 9 M", "\u2713"],
               ["B\u2084", "FinGPT-Falcon-7B-LoRA", "7 B + 9 M", "\u2713"],
               ["B\u2085", "FinGPT-Bloom-7B1-LoRA", "7 B + 9 M", "\u2713"]])

    doc.add_paragraph()

    subsection_heading(doc, "5.2", "Metrics")

    p = body_paragraph(doc, [
        {"text": "Accuracy: ", "bold": True},
        {"text": "fraction of correct predictions."}
    ])

    p = body_paragraph(doc, [
        {"text": "Weighted F1: ", "bold": True},
        {"text": "F1 averaged with class frequency weights (robust to class imbalance)."}
    ])

    subsection_heading(doc, "5.3", "Main Result")

    simple_body(doc, "Finistral-7B-LoRA improves upon the strongest FinGPT baseline by +78 F1 "
        "points and the Mistral-7B by +70 F1 points.")

    # Table 7: Test Results
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    add_run(p, "Table 7: ", bold=True, font_size=9)
    add_run(p, "Test Results", font_size=9)

    add_table(doc,
              ["Model", "Accuracy", "Weighted F1"],
              [[("Finistral-7B-LoRA (ours)", True), ("0.9956", True), ("0.9956", True)],
               ["Mistral-7B-Base", "0.4125", "0.2967"],
               ["FinGPT-mt-Llama2-7B-LoRA", "0.1564", "0.0789"],
               ["FinGPT-Llama-3-8B-LoRA", "0.2310", "0.2112"],
               ["FinGPT-Falcon-7B-LoRA", "0.2102", "0.1713"],
               ["FinGPT-Bloom-7B1-LoRA", "0.1312", "0.0310"]])

    doc.add_paragraph()

    simple_body(doc, "Only ten of 2,264 samples are misclassified; errors mostly involve sarcastic "
        "headlines or mixed-sentiment sentences.")

    subsection_heading(doc, "5.4", "Hardware & Software")

    simple_body(doc, "2 \u00d7 NVIDIA A100 40 GB on Indiana University\u2019s Big Red 200.")
    simple_body(doc, "PyTorch 2.2, Transformers 4.39, PEFT 0.10, bitsandbytes 0.43.")
    simple_body(doc, "Evaluation script performs tokenised generation followed by regex label mapping.")

    # ══════════════════════════════════════════════════════════════════════════
    # 6. RESULT ANALYSIS & DISCUSSION
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, 6, "Result Analysis & Discussion")

    subsection_heading(doc, "6.1", "Quantitative Gains")

    p = body_paragraph(doc, [
        {"text": "Finistral-7B-LoRA delivers "},
        {"text": "state-of-the-art performance", "bold": True},
        {"text": " on "},
        {"text": "Financial PhraseBank dataset", "italic": True},
        {"text": ":"}
    ])

    for bullet in ["\u2022  +69 F1 over the zero-shot Mistral-7B backbone",
                   "\u2022  +78\u201396 F1 over every FinGPT LoRA baseline (Llama2, Llama3-8B, Falcon-7B, Bloom-7B1)",
                   "\u2022  \u2248 +3\u20136 accuracy points beyond published FinBERT scores"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        add_run(p, bullet, font_size=10.5)

    simple_body(doc, "The adapter reduces test errors from > 1,400 (Mistral-7B base) to 10.")
    simple_body(doc, "All three classes are handled evenly: Negative (98.7% F1), Neutral (99.4%), "
        "Positive (99.3%).")
    simple_body(doc, "The confusion matrix (Fig. 5) shows no systematic bias toward the majority "
        "neutral class.")

    # Figures 3, 4, 5
    add_figure(doc, os.path.join(FIGURES_DIR, "image3.png"), 4.0,
               "Finistral-7B-LoRA vs Mistral-7B-Base on Financial PhraseBank (sentences_allagree)", 3)
    add_figure(doc, os.path.join(FIGURES_DIR, "image4.png"), 4.0,
               "Finistral-7B-LoRA vs FinGPT Baseline LoRA on Financial PhraseBank", 4)
    add_figure(doc, os.path.join(FIGURES_DIR, "image5.png"), 3.6,
               "Confusion matrix \u2013 Finistral-7B-LoRA", 5)

    subsection_heading(doc, "6.2", "Training Dynamics Insight")

    simple_body(doc, "Validation loss reached its minimum (0.1008) at epoch 2 and then rose, a "
        "classic sign of incipient overfitting. Selecting the epoch-2 checkpoint therefore maximises "
        "generalisation.")

    subsection_heading(doc, "6.3", "Efficiency vs. Quality Tradeoffs")

    # Table 8
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    add_run(p, "Table 8: ", bold=True, font_size=9)
    add_run(p, "Efficiency vs. Quality Tradeoffs", font_size=9)

    add_table(doc,
              ["Dimension", "Finistral-7B-LoRA", "Best FinGPT adapter"],
              [["Trainable params", "9 M", "9 M"],
               ["Adapter size", "83.9 MB (fp16)", "83\u201387 MB"],
               ["Train time (2\u00d7A100)", "\u2248 120 min", "80\u2013120 min"],
               ["Weighted F1", "0.996", "0.211 (Llama-3-8B)"]])

    doc.add_paragraph()

    subsection_heading(doc, "6.4", "Deployment Footprint")

    p = body_paragraph(doc, [
        {"text": "GPU inference", "bold": True},
        {"text": " \u2013 10 ms per sentence (A100, bf16)."}
    ])

    p = body_paragraph(doc, [
        {"text": "CPU inference", "bold": True},
        {"text": " \u2013 78 ms (M1 Pro, 4-bit GGML)."}
    ])

    simple_body(doc, "Both meet sub-second latency targets for real-time news feeds.")

    subsection_heading(doc, "6.5", "Limitations & Ethical Considerations")

    for label, text in [
        ("Label noise: ", "FinGPT relies on distant supervision; some training sentences may be mislabelled, potentially propagating bias."),
        ("Domain scope: ", "The model is English-only and tuned on equity-centric text; performance on commodities or multilingual filings is untested."),
        ("Market impact: ", "Releasing highly accurate sentiment models could amplify herding behaviour if widely adopted."),
        ("Overconfidence: ", "Near-perfect headline accuracy does not imply robustness to longer documents or adversarial phrasing; calibration techniques should be explored."),
    ]:
        body_paragraph(doc, [{"text": label, "bold": True}, {"text": text}])

    subsection_heading(doc, "6.6", "Key Takeaways")

    simple_body(doc, "LoRA + domain data + prompt engineering \u21d2 near-perfect accuracy with "
        "negligible compute.")
    simple_body(doc, "Training beyond the validation-loss trough yields no benefit and can harm "
        "generalisation.")
    simple_body(doc, "Remaining errors cluster around sarcasm and conditional language\u2014prime "
        "targets for future dataset curation or multi-label extensions.")

    # ══════════════════════════════════════════════════════════════════════════
    # 7. CONCLUSION
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, 7, "Conclusion")

    p = body_paragraph(doc, [
        {"text": "Finistral-7B-LoRA demonstrates that "},
        {"text": "state-of-the-art financial sentiment accuracy (99.6% / 0.996 F1) can be reached "
         "with only 0.2% trainable parameters, one-node GPU time, and an 84 MB adapter file", "bold": True},
        {"text": ". By combining the compute-efficient Mistral-7B backbone with Low-Rank Adaptation "
         "and the FinGPT sentiment corpus, we surpass both full-model FinBERT fine-tuning and a "
         "suite of FinGPT LoRA baselines while staying within academic compute budgets on Indiana "
         "University\u2019s Big Red 200 supercomputer. Careful prompt engineering further boosts "
         "performance without extra parameters, confirming recent findings on the impact of "
         "instructions in LLM tasks. The project validates the broader PEFT vision that small, "
         "targeted updates can rival or exceed heavy full fine-tuning approaches."}
    ])

    # ══════════════════════════════════════════════════════════════════════════
    # 8. FUTURE WORK
    # ══════════════════════════════════════════════════════════════════════════

    section_heading(doc, 8, "Future Work")

    for label, text in [
        ("Streaming EDGAR & real-time news: ",
         "Implement continual learning to keep the model current as market language evolves, "
         "aligning with FinNLP shared-task goals."),
        ("Human-in-the-loop curation: ",
         "Use active learning loops to clean noisy FinGPT labels and focus annotation on sarcasm "
         "or mixed-tone edge cases revealed in our error analysis."),
        ("4-bit & QLoRA fusion: ",
         "Combine LoRA with GGML/QLoRA [6] quantisation for sub-100 MB end-to-end models "
         "deployable on laptops and edge devices."),
        ("GRPO-based reasoning integration: ",
         "Apply Group Relative Policy Optimization (GRPO) [12] to inject chain-of-thought "
         "reasoning into Finistral-7B-LoRA, enabling the model to articulate its sentiment "
         "rationale rather than returning a bare label. GRPO eliminates the need for a separate "
         "value model, making reinforcement-learning-based alignment feasible within academic "
         "compute budgets."),
        ("On-device small language models: ",
         "Distil the reasoning-augmented adapter into sub-3B parameter small language models (SLMs) "
         "optimised for on-device inference on consumer hardware such as iPhone 15, Android "
         "smartphones, and low-end laptops. Combining GRPO-trained reasoning with aggressive "
         "quantisation (4-bit/2-bit) and architecture pruning could yield models that run entirely "
         "on-device with sub-200 ms latency, eliminating cloud dependency and preserving user "
         "privacy for personal finance applications."),
    ]:
        body_paragraph(doc, [{"text": label, "bold": True}, {"text": text}])

    # ══════════════════════════════════════════════════════════════════════════
    # BACK MATTER
    # ══════════════════════════════════════════════════════════════════════════

    unnumbered_heading(doc, "Acknowledgments")
    simple_body(doc, "The authors acknowledge Indiana University\u2019s Big Red 200 supercomputing "
        "facility for providing the GPU resources used in this study.")

    # CRediT
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, "CRediT Author Contribution Statement", bold=True, font_size=10.5)

    p = body_paragraph(doc, [
        {"text": "Zainab Mirza: ", "bold": True},
        {"text": "Supervision; Writing \u2013 Review & Editing. "},
        {"text": "Ayan Javeed Shaikh: ", "bold": True},
        {"text": "Conceptualization; Methodology; Software; Validation; Formal Analysis; "
         "Investigation; Data Curation; Writing \u2013 Original Draft; Writing \u2013 Review & "
         "Editing; Visualization. "},
        {"text": "Fazal Jalil Parkar: ", "bold": True},
        {"text": "Writing \u2013 Review & Editing."}
    ])

    unnumbered_heading(doc, "Conflict of Interest")
    simple_body(doc, "There is no conflict of interest.")

    unnumbered_heading(doc, "Supporting Information")
    simple_body(doc, "The trained adapter weights, tokeniser files, inference scripts, and Gradio "
        "GUI are publicly available at https://huggingface.co/Ayansk11/Finistral-7B_lora. The "
        "complete source code, training scripts, evaluation notebooks, and this manuscript\u2019s "
        "LaTeX source are maintained at https://github.com/ayansk11/FinistralAI_code.")

    unnumbered_heading(doc, "Use of Artificial Intelligence (AI)-Assisted Technology for Manuscript Preparation")
    simple_body(doc, "The authors confirm that there was no use of artificial intelligence "
        "(AI)-assisted technology for assisting in the writing or editing of the manuscript and no "
        "images were manipulated using AI.")

    # ══════════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ══════════════════════════════════════════════════════════════════════════

    unnumbered_heading(doc, "References")

    refs = [
        "Loughran, T., McDonald, B., When Is a Liability Not a Liability? Textual Analysis, Dictionaries, and 10-Ks, The Journal of Finance, 2011, 66(1), 35\u201365, doi: 10.1111/j.1540-6261.2010.01625.x.",
        "Malo, P., Sinha, A., Korhonen, P., Wallenius, J., Takala, P., Good Debt or Bad Debt: Detecting Semantic Orientations in Economic Texts, Journal of the Association for Information Science and Technology, 2014, 65(4), 782\u2013796, doi: 10.1002/asi.23062.",
        "Devlin, J., Chang, M.W., Lee, K., Toutanova, K., BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding, Proceedings of the 2019 Conference of the North American Chapter of the Association for Computational Linguistics (NAACL), 2019, 4171\u20134186. https://aclanthology.org/N19-1423/",
        "Araci, D., FinBERT: Financial Sentiment Analysis with Pre-trained Language Models, arXiv preprint, 2019, arXiv:1908.10063. https://arxiv.org/abs/1908.10063",
        "Hu, E.J., Shen, Y., Wallis, P., Allen-Zhu, Z., Li, Y., Wang, S., Wang, L., Chen, W., LoRA: Low-Rank Adaptation of Large Language Models, International Conference on Learning Representations (ICLR), 2022. https://arxiv.org/abs/2106.09685",
        "Dettmers, T., Pagnoni, A., Holtzman, A., Zettlemoyer, L., QLoRA: Efficient Finetuning of Quantized LLMs, Advances in Neural Information Processing Systems (NeurIPS), 2023, arXiv:2305.14314. https://arxiv.org/abs/2305.14314",
        "Han, Z., Gao, C., Liu, J., Zhang, J., Zhang, S.Q., Parameter-Efficient Fine-Tuning for Large Models: A Comprehensive Survey, arXiv preprint, 2024, arXiv:2403.14608. https://arxiv.org/abs/2403.14608",
        "Yang, H., Liu, X.Y., Wang, C.D., FinGPT: Open-Source Financial Large Language Models, arXiv preprint, 2023, arXiv:2306.06031. https://arxiv.org/abs/2306.06031",
        "Zhang, B., Yang, H., Liu, X.Y., Instruct-FinGPT: Financial Sentiment Analysis by Instruction Tuning of General-Purpose Large Language Models, FinLLM Symposium at IJCAI, 2023, arXiv:2306.12659. https://arxiv.org/abs/2306.12659",
        "Chen, C., Huang, H.H., Chen, H.H., et al., Proceedings of the Sixth Workshop on Financial Technology and Natural Language Processing (FinNLP 2023), Association for Computational Linguistics, 2023. https://aclanthology.org/2023.finnlp-2.0.pdf",
        "Jiang, A.Q., Sablayrolles, A., Mensch, A., Bamford, C., Chaplot, D.S., de las Casas, D., et al., Mistral 7B, arXiv preprint, 2023, arXiv:2310.06825. https://arxiv.org/abs/2310.06825",
        "Shao, Z., Wang, P., Zhu, Q., Xu, R., Song, J., Zhang, M., et al., DeepSeekMath: Pushing the Limits of Mathematical Reasoning in Open Language Models, arXiv preprint, 2024, arXiv:2402.03300. https://arxiv.org/abs/2402.03300",
        "Han, D., Han, M., Unsloth: Fine-tuning LLMs 2x Faster with 70% Less Memory, 2023. https://unsloth.ai/docs",
    ]
    for i, ref in enumerate(refs, 1):
        add_reference(doc, i, ref)

    # ══════════════════════════════════════════════════════════════════════════
    # PUBLISHER NOTE & OPEN ACCESS
    # ══════════════════════════════════════════════════════════════════════════

    doc.add_paragraph()

    p = body_paragraph(doc, [
        {"text": "Publisher Note: ", "bold": True},
        {"text": "The views, statements, and data in all publications solely belong to the authors "
         "and contributors. GR Scholastic is not responsible for any injury resulting from the "
         "ideas, methods, or products mentioned. GR Scholastic remains neutral regarding "
         "jurisdictional claims in published maps and institutional affiliations."}
    ], space_before=Pt(12))

    bold_body(doc, "Open Access")

    simple_body(doc, "This article is licensed under a Creative Commons "
        "Attribution-NonCommercial 4.0 International License, which permits the non-commercial use, "
        "sharing, adaptation, distribution and reproduction in any medium or format, as long as "
        "appropriate credit to the original author(s) and the source is given by providing a link "
        "to the Creative Commons License and changes need to be indicated if there are any.")

    simple_body(doc, "To view a copy of this License, visit: "
        "https://creativecommons.org/licenses/by-nc/4.0/")

    simple_body(doc, "\u00a9 The Authors 2025")

    # ── Save ──
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
